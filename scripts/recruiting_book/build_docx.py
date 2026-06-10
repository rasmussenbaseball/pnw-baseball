"""
Build a native Word .docx version of the Pacific Northwest College Baseball
Program Guide.

The PDF renders perfectly in any PDF viewer (Drive's preview, Acrobat,
Preview, Chrome), but Google Docs' PDF-to-Doc converter merges adjacent
table rows of the same style into one paragraph, which mashes the TOC
together. The fix is to ship a native .docx so users who want a Word-style
view get one without going through the lossy PDF→Doc conversion.

This script reads scripts/recruiting_book/book_data.json (built by
gather_book_data.py) and writes
scripts/recruiting_book/PNW_College_Baseball_Guide_2026.docx with the same
structure as the PDF: cover, about, table of contents, then five level
chapters with one section per program.

Run after gather_book_data.py:
    python3 scripts/recruiting_book/build_docx.py
"""
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor, Inches

HERE = Path(__file__).resolve().parent
_RAW = json.loads((HERE / "book_data.json").read_text())
DATA = _RAW["teams"] if isinstance(_RAW, dict) and "teams" in _RAW else _RAW
LEVEL_NORMS = (_RAW.get("level_norms") if isinstance(_RAW, dict) else None) or {}
OUT = HERE / "PNW_College_Baseball_Guide_2026.docx"

# Same level colors / titles as the PDF.
LEVELS = [
    {"key": "D1",   "title": "Division I",   "subtitle": "NCAA Division I",                                       "color": "14365c"},
    {"key": "D2",   "title": "Division II",  "subtitle": "NCAA Division II",                                      "color": "0d7d7d"},
    {"key": "D3",   "title": "Division III", "subtitle": "NCAA Division III",                                     "color": "2f7d4f"},
    {"key": "NAIA", "title": "NAIA",         "subtitle": "National Association of Intercollegiate Athletics",     "color": "b07d12"},
    {"key": "NWAC", "title": "NWAC",         "subtitle": "Northwest Athletic Conference (Junior College)",        "color": "9c1f2e"},
]
LEVEL_COLOR_HEX = {lv["key"]: lv["color"] for lv in LEVELS}

PRIMERS = {
    "D1": "Division I is the highest level of college baseball. After the House v. NCAA settlement that took effect for 2025-26, D1 baseball is now under a 34-player roster cap, and every player on that roster is eligible for a full athletic scholarship. Programs operate on large budgets and recruit nationally. The Pacific Northwest fields seven D1 programs across three states.",
    "D2": "Division II programs fund the equivalent of nine athletic scholarships. Rosters are more regional than D1, with development as the central theme. The Great Northwest Athletic Conference is the home for D2 baseball in this region — a compact, travel-friendly league.",
    "D3": "Division III baseball has no athletic scholarships. Players are recruited but pay through need-based and merit-based academic aid. Fit, major, and the financial package drive every decision. The Northwest Conference anchors D3 baseball in this region.",
    "NAIA": "The NAIA allows up to 12 scholarship equivalencies, more than D1 or D2, with friendlier eligibility and transfer rules. The Cascade Collegiate Conference anchors NAIA baseball across Oregon and Idaho.",
    "NWAC": "The NWAC is the Northwest's junior-college level. Athletic aid is allowed but capped at 65 percent of in-state tuition under conference rules. Combined with low community-college tuition, the total cost of attendance can still come in lower than four-year options. The league spans Washington, Oregon, Idaho, and British Columbia across four geographic conferences.",
}


# ───── helpers ──────────────────────────────────────────────────────────────

def fmt_count(v):
    if v is None or v == "":
        return ""
    try:
        n = int(re.sub(r"\D", "", str(v)))
        return f"{n:,}"
    except Exception:
        return str(v)


def fmt_avg(v):
    try:
        s = "%.3f" % float(v)
        return s.lstrip("0") if s.startswith("0") else s
    except Exception:
        return "—"


def fmt_rate(v, places=2):
    if v is None:
        return "—"
    try:
        return ("%." + str(places) + "f") % float(v)
    except Exception:
        return "—"


def fmt_int(v):
    try:
        return f"{int(v):,}"
    except Exception:
        return "—"


def fmt_ip(v):
    try:
        return "%.1f" % float(v)
    except Exception:
        return "—"


def rec_str(s):
    if not s:
        return "—"
    w, l, t = s.get("wins") or 0, s.get("losses") or 0, s.get("ties") or 0
    return f"{w}-{l}-{t}" if t else f"{w}-{l}"


def ordinal(n):
    if n is None:
        return ""
    n = int(n)
    if 10 <= n % 100 <= 20:
        suf = "th"
    else:
        suf = {1: "st", 2: "nd", 3: "rd"}.get(n % 10, "th")
    return f"{n}{suf}"


def g(profile, key, default=""):
    v = profile.get(key)
    return "" if v is None else (str(v).strip() or default)


def is_tbd(s):
    return s and re.search(r"\btbd\b|new coach", s, re.I) is not None


def scholarship_line(team):
    div = team["division"]
    raw = g(team["profile"], "scholarshipInfo")
    if div == "D1":
        return "Up to 34 (roster cap, full scholarships allowed)"
    if div == "D2":
        return raw or "Up to 9.0 scholarship equivalencies"
    if div == "D3":
        return "No athletic scholarships (academic and need-based aid only)"
    if div == "NAIA":
        return raw or "Up to 12 scholarship equivalencies"
    if div == "NWAC":
        return "Athletic aid available, capped at 65% of in-state tuition (NWAC rule)"
    return raw


# ───── docx helpers ─────────────────────────────────────────────────────────

def set_cell_bg(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def add_heading(doc, text, level=1, color_hex=None):
    h = doc.add_heading(text, level=level)
    if color_hex and h.runs:
        for r in h.runs:
            r.font.color.rgb = RGBColor.from_string(color_hex)
    return h


def add_para(doc, text, bold=False, italic=False, size=None, color_hex=None, align=None):
    p = doc.add_paragraph()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size is not None:
        run.font.size = Pt(size)
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    return p


def add_kv_table(doc, rows, color_hex):
    """Two-column label/value table. Label cells get an accent bar fill."""
    if not rows:
        return None
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light Grid"
    table.autofit = False
    for i, (lbl, val) in enumerate(rows):
        c0 = table.cell(i, 0)
        c1 = table.cell(i, 1)
        c0.text = ""
        c1.text = ""
        # Label cell.
        rp = c0.paragraphs[0].add_run(lbl)
        rp.bold = True
        rp.font.size = Pt(9)
        rp.font.color.rgb = RGBColor.from_string("5f5f5f")
        # Value cell.
        rv = c1.paragraphs[0].add_run(val)
        rv.font.size = Pt(10)
    # Set widths.
    for row in table.rows:
        row.cells[0].width = Inches(2.0)
        row.cells[1].width = Inches(4.5)
    return table


def add_stats_grid(doc, title, pairs, color_hex):
    """Compact 2-row grid: top row = header (one merged cell), then label/value pairs."""
    if not pairs:
        return None
    p = doc.add_paragraph()
    r = p.add_run(title.upper())
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string("5f5f5f")
    # The actual data table: each row is "Label  Value  Label  Value".
    n_rows = (len(pairs) + 1) // 2
    table = doc.add_table(rows=n_rows, cols=4)
    table.style = "Light List"
    for i in range(n_rows):
        a = pairs[2 * i]
        b = pairs[2 * i + 1] if 2 * i + 1 < len(pairs) else ("", "")
        for col, (lbl, val) in enumerate([a, b]):
            c_lbl = table.cell(i, col * 2)
            c_val = table.cell(i, col * 2 + 1)
            c_lbl.text = ""
            c_val.text = ""
            rl = c_lbl.paragraphs[0].add_run(lbl)
            rl.bold = True
            rl.font.size = Pt(8)
            rl.font.color.rgb = RGBColor.from_string("8a8a8a")
            rv = c_val.paragraphs[0].add_run(str(val))
            rv.font.size = Pt(10)
    return table


# ───── prose ────────────────────────────────────────────────────────────────

def program_intro_text(team):
    p = team["profile"]
    name = p.get("teamName") or team["school_name"]
    conf = team.get("conference") or ""
    loc = ", ".join(x for x in [team.get("city"), team.get("state")] if x)
    enr = fmt_count(g(p, "enrollment"))
    stype = g(p, "schoolType")
    setting = g(p, "campusSetting")

    para1_bits = [f"{name} play out of {team['school_name']}"]
    if loc:
        para1_bits[-1] += f" in {loc}"
    if conf:
        para1_bits[-1] += f", and compete in the {conf}"
    para1_bits[-1] += "."

    desc_parts = []
    if stype:
        inst = stype.lower()
        desc_parts.append(f"a {inst} school" if ("public" in inst or "private" in inst) else f"a {inst} institution")
    if enr:
        if desc_parts:
            desc_parts[-1] += f" with roughly {enr} undergraduates"
        else:
            desc_parts.append(f"a campus of roughly {enr} undergraduates")
    if setting and desc_parts:
        desc_parts[-1] += f" in a {setting.lower()} setting"
    if desc_parts:
        para1_bits.append(f"It is {desc_parts[0]}.")

    accept = g(p, "acceptance")
    grad = g(p, "gradRate")
    majors = g(p, "topMajors")
    acad_bits = []
    if accept:
        acad_bits.append("admissions are open enrollment" if accept.lower() == "open enrollment" else f"the acceptance rate sits near {accept}")
    if grad:
        acad_bits.append(f"about {grad} of students graduate")
    if acad_bits:
        para1_bits.append("Academically, " + " and ".join(acad_bits) + ".")
    if majors:
        para1_bits.append(f"Common areas of study include {majors}.")

    return " ".join(para1_bits)


def season_recap_text(team):
    s = team.get("season_2026")
    if not s:
        return "The 2026 season's complete statistical record was not available for this program."
    name = team["short_name"]
    w, l = s.get("wins") or 0, s.get("losses") or 0
    cw, cl = s.get("conference_wins") or 0, s.get("conference_losses") or 0
    rd = s.get("run_differential")
    rs, ra = s.get("runs_scored"), s.get("runs_allowed")
    avg = s.get("team_batting_avg")
    ops = s.get("team_ops")
    era = s.get("team_era")
    cs = team.get("conf_standings") or {}
    place = None
    if cs.get("standings"):
        place = next((r["place"] for r in cs["standings"] if r["id"] == team["team_id"]), None)
    conf_size = len(cs.get("standings") or [])

    bits = [f"{name} finished {w}-{l} in 2026 and went {cw}-{cl} against {team.get('conference','conference')} opponents."]
    if place and conf_size:
        bits.append(f"That landed them {ordinal(place)} out of {conf_size} teams.")
    if rd is not None and (rs or ra):
        verb = "outscored opponents" if rd >= 0 else "were outscored"
        bits.append(f"They {verb} by {abs(int(rd))} runs ({int(rs or 0)} scored, {int(ra or 0)} allowed).")
    if avg or ops or era:
        sub = []
        if avg: sub.append(f"a {fmt_avg(avg)} team average")
        if ops: sub.append(f"{fmt_avg(ops)} OPS")
        if era: sub.append(f"and a {fmt_rate(era)} ERA")
        if sub:
            bits.append("The line: " + ", ".join(sub) + ".")
    th = team.get("top_hitter")
    tp = team.get("top_pitcher")
    if th and th.get("name"):
        slash = ""
        if th.get("batting_avg") and th.get("on_base_pct") and th.get("slugging_pct"):
            slash = f", slashing {fmt_avg(th['batting_avg'])}/{fmt_avg(th['on_base_pct'])}/{fmt_avg(th['slugging_pct'])}"
        hr = f", {int(th['home_runs'])} HR" if th.get("home_runs") else ""
        rbi = f", {int(th['rbi'])} RBI" if th.get("rbi") else ""
        bits.append(f"{th['name']} was the standout bat{slash}{hr}{rbi}.")
    if tp and tp.get("name"):
        era_p = f"{fmt_rate(tp['era'])} ERA" if tp.get("era") is not None else ""
        ip_p = f", {fmt_ip(tp['innings_pitched'])} IP" if tp.get("innings_pitched") is not None else ""
        k_p = f", {int(tp['strikeouts'])} K" if tp.get("strikeouts") else ""
        bits.append(f"On the mound, {tp['name']} led the staff with a {era_p}{ip_p}{k_p}.")
    return " ".join(bits)


# ───── chapter sections ─────────────────────────────────────────────────────

def add_team_section(doc, team):
    color = LEVEL_COLOR_HEX[team["division"]]
    name = team["profile"].get("teamName") or team["school_name"]
    # Heading 1: school name (acts as a TOC entry in Word's native TOC).
    h = doc.add_heading(team["school_name"], level=1)
    if h.runs:
        for r in h.runs:
            r.font.color.rgb = RGBColor.from_string(color)
    # Sub-heading: team name / conference / location.
    sub_bits = [team.get("mascot"), team.get("conference"),
                ", ".join(x for x in [team.get("city"), team.get("state")] if x)]
    add_para(doc, "  ·  ".join(x for x in sub_bits if x),
             italic=True, size=10, color_hex="5f5f5f")

    # Snapshot stat strip.
    s = team.get("season_2026") or {}
    finish = ""
    cs = team.get("conf_standings")
    if cs and cs.get("standings"):
        place = next((r["place"] for r in cs["standings"] if r["id"] == team["team_id"]), None)
        if place:
            finish = f"{ordinal(place)} of {len(cs['standings'])}"
    tiles = [
        ("2026 RECORD", rec_str(s)),
        ("CONFERENCE", f"{s.get('conference_wins',0)}-{s.get('conference_losses',0)}" if s else "—"),
        ("RUN DIFF", (("+" if (s.get('run_differential') or 0) >= 0 else "") + str(int(s['run_differential']))) if s.get('run_differential') is not None else "—"),
        ("TEAM AVG", fmt_avg(s.get("team_batting_avg")) if s.get("team_batting_avg") else "—"),
        ("TEAM ERA", fmt_rate(s.get("team_era")) if s.get("team_era") else "—"),
        ("CONF FINISH" if finish else "ENROLLMENT", finish or (fmt_count(g(team["profile"], "enrollment")) or "—")),
    ]
    table = doc.add_table(rows=2, cols=len(tiles))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for col, (lbl, val) in enumerate(tiles):
        c_lbl = table.cell(0, col)
        c_val = table.cell(1, col)
        set_cell_bg(c_lbl, color)
        c_lbl.text = ""
        c_val.text = ""
        r_lbl = c_lbl.paragraphs[0].add_run(lbl)
        r_lbl.bold = True
        r_lbl.font.size = Pt(8)
        r_lbl.font.color.rgb = RGBColor.from_string("FFFFFF")
        c_lbl.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_val = c_val.paragraphs[0].add_run(val)
        r_val.bold = True
        r_val.font.size = Pt(12)
        c_val.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "The Program", level=2, color_hex=color)
    add_para(doc, program_intro_text(team), size=11)
    add_heading(doc, "The 2026 Season", level=2, color_hex=color)
    add_para(doc, season_recap_text(team), size=11)

    # Top hitters.
    hitters = team.get("top_hitters") or []
    if hitters:
        add_heading(doc, "Top Hitters by WAR (min. 40 PA)", level=2, color_hex=color)
        cols = ["Player", "Pos", "Yr", "PA", "AVG", "OBP", "SLG", "OPS", "HR", "RBI", "SB", "wRC+", "WAR"]
        tbl = doc.add_table(rows=1 + len(hitters[:5]), cols=len(cols))
        tbl.style = "Light Grid"
        for col, h in enumerate(cols):
            c = tbl.cell(0, col)
            set_cell_bg(c, color)
            c.text = ""
            r = c.paragraphs[0].add_run(h)
            r.bold = True
            r.font.size = Pt(8)
            r.font.color.rgb = RGBColor.from_string("FFFFFF")
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, h in enumerate(hitters[:5]):
            row = tbl.rows[i + 1]
            vals = [
                h.get("name") or "—",
                h.get("position") or "—",
                (h.get("year_in_school") or "—")[:4],
                fmt_int(h.get("plate_appearances")),
                fmt_avg(h.get("batting_avg")),
                fmt_avg(h.get("on_base_pct")),
                fmt_avg(h.get("slugging_pct")),
                fmt_avg(h.get("ops")),
                fmt_int(h.get("home_runs")),
                fmt_int(h.get("rbi")),
                fmt_int(h.get("stolen_bases")),
                fmt_int(h.get("wrc_plus")) if h.get("wrc_plus") else "—",
                ("%.2f" % float(h["offensive_war"])) if h.get("offensive_war") is not None else "—",
            ]
            for col, v in enumerate(vals):
                cc = row.cells[col]
                cc.text = ""
                rr = cc.paragraphs[0].add_run(str(v))
                rr.font.size = Pt(8)
                if col == 0:
                    rr.bold = True
                cc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    pitchers = team.get("top_pitchers") or []
    if pitchers:
        add_heading(doc, "Top Pitchers by WAR (min. 15 IP)", level=2, color_hex=color)
        cols = ["Player", "Yr", "G", "GS", "W-L", "SV", "IP", "ERA", "WHIP", "K", "K/9", "FIP", "WAR"]
        tbl = doc.add_table(rows=1 + len(pitchers[:5]), cols=len(cols))
        tbl.style = "Light Grid"
        for col, h in enumerate(cols):
            c = tbl.cell(0, col)
            set_cell_bg(c, color)
            c.text = ""
            r = c.paragraphs[0].add_run(h)
            r.bold = True
            r.font.size = Pt(8)
            r.font.color.rgb = RGBColor.from_string("FFFFFF")
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, p in enumerate(pitchers[:5]):
            row = tbl.rows[i + 1]
            wl = f"{p.get('wins') or 0}-{p.get('losses') or 0}"
            vals = [
                p.get("name") or "—",
                (p.get("year_in_school") or "—")[:4],
                fmt_int(p.get("games")),
                fmt_int(p.get("games_started")),
                wl,
                fmt_int(p.get("saves")) if p.get("saves") else "0",
                fmt_ip(p.get("innings_pitched")),
                fmt_rate(p.get("era")),
                fmt_rate(p.get("whip")),
                fmt_int(p.get("strikeouts")),
                fmt_rate(p.get("k_per_9"), 1),
                fmt_rate(p.get("fip")),
                ("%.2f" % float(p["pitching_war"])) if p.get("pitching_war") is not None else "—",
            ]
            for col, v in enumerate(vals):
                cc = row.cells[col]
                cc.text = ""
                rr = cc.paragraphs[0].add_run(str(v))
                rr.font.size = Pt(8)
                if col == 0:
                    rr.bold = True
                cc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Coach.
    p = team["profile"]
    coach = g(p, "coach")
    if coach and not is_tbd(coach):
        add_heading(doc, "Head Coach", level=2, color_hex=color)
        meta_bits = []
        if g(p, "coachYears"):
            meta_bits.append(g(p, "coachYears"))
        if g(p, "coachAlma"):
            meta_bits.append(f"{g(p,'coachAlma')} alum")
        if g(p, "careerRecord") and any(ch.isdigit() for ch in g(p, "careerRecord")):
            meta_bits.append(f"Career: {g(p,'careerRecord')}")
        head_line = f"{coach}  ·  " + "  ·  ".join(meta_bits) if meta_bits else coach
        add_para(doc, head_line, bold=True, color_hex=color, size=11)
        if g(p, "prevStops"):
            add_para(doc, f"Coaching path: {g(p,'prevStops')}", size=10)
        if g(p, "coachEmail"):
            add_para(doc, f"Email: {g(p,'coachEmail')}", size=10)
        bio = g(p, "coachBio")
        if bio:
            # Trim long bios at ~2400 chars on a sentence break.
            bio_clean = re.sub(r"\s+", " ", bio).strip()
            if len(bio_clean) > 2400:
                cut = bio_clean[:2400]
                idx = cut.rfind(". ")
                if idx > 1400:
                    bio_clean = cut[:idx + 1].strip()
            add_para(doc, bio_clean, size=10)
        if g(p, "assistants"):
            add_para(doc, f"Coaching staff: {g(p,'assistants')}", size=10, italic=True)

    # Program facts (academics + cost + facilities + recruiting).
    add_heading(doc, "Program Facts", level=2, color_hex=color)
    fact_rows = []
    def add(lbl, val):
        if val:
            fact_rows.append((lbl, val))
    add("Stadium", " ".join(x for x in [g(p, "stadium"),
                                        (f"({g(p,'capacity')})" if g(p,'capacity') else "")] if x))
    add("School Type", g(p, "schoolType"))
    add("Enrollment", fmt_count(g(p, "enrollment")))
    add("Student / Faculty", g(p, "sfr"))
    add("Acceptance Rate", g(p, "acceptance"))
    add("Top Majors", g(p, "topMajors"))
    add("Graduation Rate", g(p, "gradRate"))
    add("Financial Aid", g(p, "financialAidPct"))
    add("In-State Tuition", g(p, "inStateTuition"))
    add("Out-of-State Tuition", g(p, "outStateTuition"))
    add("Room & Board", g(p, "roomBoard"))
    add("Scholarships", scholarship_line(team))
    add("Roster Size", g(p, "rosterSize"))
    add("2026 Seniors", g(p, "gradSeniors"))
    add("Nearest Airport", g(p, "nearestAirport"))
    add("Distance", g(p, "distanceFromCity"))
    add("Campus Setting", g(p, "campusSetting"))
    add("Campus Safety", g(p, "campusSafety"))
    add("Athletics Site", g(p, "athleticsWebsite"))
    add("Roster Page", g(p, "baseballRosterUrl"))
    add("School Site", g(p, "schoolWebsite"))
    add_kv_table(doc, fact_rows, color)

    # Pro alumni.
    alumni = team.get("pro_alumni") or []
    if alumni:
        add_heading(doc, "Currently in Pro Baseball", level=2, color_hex=color)
        mlb_count = sum(1 for a in alumni if a.get("level") == "MLB")
        by_lvl = {}
        for a in alumni:
            if a.get("level") and a["level"] != "MLB":
                by_lvl[a["level"]] = by_lvl.get(a["level"], 0) + 1
        parts = []
        if mlb_count:
            parts.append(f"{mlb_count} MLB")
        for lvl in ("AAA", "AA", "A+", "A", "Rk"):
            if by_lvl.get(lvl):
                parts.append(f"{by_lvl[lvl]} {lvl}")
        add_para(doc, f"{len(alumni)} total alumni currently in pro baseball  ·  " + "  ·  ".join(parts), size=10)
        for a in alumni[:20]:
            line = a.get("name") or ""
            if a.get("level"):
                line += f"  ·  {a['level']}"
            if a.get("current_team") or a.get("affiliate"):
                line += f", {a.get('current_team') or a.get('affiliate')}"
            add_para(doc, line, size=10)
        if len(alumni) > 20:
            add_para(doc, f"… and {len(alumni) - 20} more in the system.", size=10, italic=True)

    # Recruiting box.
    add_heading(doc, "How to Get on the Radar", level=2, color_hex=color)
    rec_bits = []
    if coach and not is_tbd(coach):
        rec_bits.append(f"Address communication to {coach}.")
    if g(p, "coachEmail"):
        rec_bits.append(f"Email: {g(p,'coachEmail')}.")
    if g(p, "recruitQuestionnaireUrl"):
        rec_bits.append(f"Recruiting questionnaire: {g(p,'recruitQuestionnaireUrl')}.")
    if scholarship_line(team):
        rec_bits.append(f"Athletic aid: {scholarship_line(team)}.")
    if g(p, "baseballRosterUrl"):
        rec_bits.append(f"Roster page: {g(p,'baseballRosterUrl')}.")
    add_para(doc, " ".join(rec_bits), size=11, bold=False)

    # Page break between teams.
    doc.add_page_break()


def build():
    doc = Document()
    # Page size: Letter, normal margins.
    sec = doc.sections[0]
    sec.page_height = Inches(11)
    sec.page_width = Inches(8.5)
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)

    # Cover.
    add_para(doc, "PACIFIC NORTHWEST", bold=True, size=13, align="center", color_hex="5f5f5f")
    add_para(doc, "College Baseball", bold=True, size=44, align="center")
    add_para(doc, "Program Guide", bold=True, size=44, align="center", color_hex="14365c")
    add_para(doc, "")
    add_para(doc, f"{len(DATA)} Programs  ·  Five Levels  ·  The 2026 Season",
             size=12, align="center", color_hex="5f5f5f")
    add_para(doc, "")
    add_para(doc, "A complete recruit and fan reference, built from the NWBB Stats database.",
             italic=True, size=10, align="center", color_hex="8a8a8a")
    doc.add_page_break()

    # About.
    add_heading(doc, "About This Guide", level=1, color_hex="14365c")
    add_para(doc,
             f"This guide profiles all {len(DATA)} four-year and junior-college baseball programs in the "
             "Pacific Northwest, organized by competitive level. It is built for recruits, families, "
             "coaches, and fans who want a single, deep reference to the region's college baseball "
             "landscape, from the biggest D1 programs to the smallest NWAC schools.",
             size=11)
    add_para(doc,
             "Records and statistics are drawn from the NWBB Stats season database; program details "
             "come from original research compiled by the NWBB Stats team. Where a field is missing, "
             "the source data was not available at the time of publication.",
             size=11)
    doc.add_page_break()

    # Native Word Table of Contents field. Word/Google Docs build the TOC
    # from the Heading 1 / Heading 2 styles we apply to each team section.
    add_heading(doc, "Table of Contents", level=1, color_hex="14365c")
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-2" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:t")
    fldChar3.text = "Right-click and choose 'Update Field' to populate the table of contents."
    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    run._r.append(fldChar4)
    doc.add_page_break()

    # Chapters.
    for level in LEVELS:
        teams = [t for t in DATA if t["division"] == level["key"]]
        teams.sort(key=lambda t: (t.get("conference") or "", t["school_name"]))
        # Use Heading 1 for chapter titles too (same level as schools so they
        # appear in the auto-TOC).
        h = doc.add_heading(level["title"], level=1)
        if h.runs:
            for r in h.runs:
                r.font.color.rgb = RGBColor.from_string(level["color"])
        add_para(doc, f"{level['subtitle']}  ·  {len(teams)} programs",
                 size=10, color_hex="5f5f5f", italic=True)
        add_para(doc, PRIMERS[level["key"]], size=11)
        doc.add_page_break()
        for t in teams:
            add_team_section(doc, t)

    doc.save(str(OUT))
    print(f"Built {OUT}")


if __name__ == "__main__":
    build()
